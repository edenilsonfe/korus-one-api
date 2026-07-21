"""Export AI reports to PDF, DOCX, TXT and MD."""

from __future__ import annotations

import io
import re
from datetime import date

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.services.assistant.format_reply import sanitize_llm_plain_text

REPORT_TYPE_LABELS = {
    "clinico": "Relatório Clínico",
    "escolar": "Relatório Escolar",
    "pais": "Relatório para Pais",
    "evolutivo": "Relatório Evolutivo",
}

_HEADING2 = re.compile(r"^##\s+(.+)$")
_HEADING3 = re.compile(r"^###\s+(.+)$")
_BULLET = re.compile(r"^[-*]\s+(.+)$")
_TABLE_SEPARATOR = re.compile(r"^\|[\s\-:|]+\|?\s*$")
_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _format_table_row(line: str) -> str:
    trimmed = line.strip()
    if "|" not in trimmed:
        return line
    inner = trimmed.strip("|")
    cells = [cell.strip() for cell in inner.split("|") if cell.strip()]
    return " · ".join(cells) if cells else ""


def _apply_bold_reportlab(text: str) -> str:
    escaped = text.replace("&", "&amp;").replace("<", "&lt;")
    return _BOLD.sub(r"<b>\1</b>", escaped)


def _iter_markdown_lines(content: str):
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            yield ("blank", "")
            continue
        if _TABLE_SEPARATOR.match(line):
            continue
        if line.strip().startswith("|"):
            yield ("text", _format_table_row(line))
            continue
        h2 = _HEADING2.match(line)
        if h2:
            yield ("h2", h2.group(1).strip())
            continue
        h3 = _HEADING3.match(line)
        if h3:
            yield ("h3", h3.group(1).strip())
            continue
        bullet = _BULLET.match(line)
        if bullet:
            yield ("bullet", bullet.group(1).strip())
            continue
        yield ("text", line.strip())


def _add_bold_runs(paragraph, text: str) -> None:
    parts = _BOLD.split(text)
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if index % 2 == 1:
            run.bold = True


def _header_lines(report_type: str, patient_name: str, report_date: date) -> list[str]:
    type_label = REPORT_TYPE_LABELS.get(report_type, report_type)
    return [
        "KorusFono",
        type_label,
        f"Paciente: {patient_name}",
        f"Data: {report_date.isoformat()}",
        "",
    ]


def export_txt(report_type: str, patient_name: str, report_date: date, content: str) -> bytes:
    lines = _header_lines(report_type, patient_name, report_date)
    lines.append(sanitize_llm_plain_text(content))
    return "\n".join(lines).encode("utf-8")


def export_md(report_type: str, patient_name: str, report_date: date, content: str) -> bytes:
    type_label = REPORT_TYPE_LABELS.get(report_type, report_type)
    body = (
        f"# {type_label}\n\n"
        f"**Paciente:** {patient_name}  \n"
        f"**Data:** {report_date.isoformat()}\n\n"
        f"---\n\n"
        f"{content}\n"
    )
    return body.encode("utf-8")


def export_docx(report_type: str, patient_name: str, report_date: date, content: str) -> bytes:
    doc = Document()
    doc.add_heading("KorusFono", level=1)
    type_label = REPORT_TYPE_LABELS.get(report_type, report_type)
    doc.add_heading(type_label, level=2)
    doc.add_paragraph(f"Paciente: {patient_name}")
    doc.add_paragraph(f"Data: {report_date.isoformat()}")
    doc.add_paragraph("")

    for kind, value in _iter_markdown_lines(content):
        if kind == "blank":
            doc.add_paragraph("")
        elif kind == "h2":
            doc.add_heading(value, level=2)
        elif kind == "h3":
            doc.add_heading(value, level=3)
        elif kind == "bullet":
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_bold_runs(paragraph, value)
        else:
            paragraph = doc.add_paragraph()
            _add_bold_runs(paragraph, value)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_pdf(report_type: str, patient_name: str, report_date: date, content: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    type_label = REPORT_TYPE_LABELS.get(report_type, report_type)
    story = [
        Paragraph("KorusFono", styles["Title"]),
        Paragraph(type_label, styles["Heading2"]),
        Paragraph(f"Paciente: {patient_name}", styles["Normal"]),
        Paragraph(f"Data: {report_date.isoformat()}", styles["Normal"]),
        Spacer(1, 12),
    ]
    for kind, value in _iter_markdown_lines(content):
        if kind == "blank":
            story.append(Spacer(1, 6))
        elif kind == "h2":
            story.append(Paragraph(_apply_bold_reportlab(value), styles["Heading2"]))
        elif kind == "h3":
            story.append(Paragraph(_apply_bold_reportlab(value), styles["Heading3"]))
        elif kind == "bullet":
            story.append(Paragraph(f"• {_apply_bold_reportlab(value)}", styles["Normal"]))
        else:
            story.append(Paragraph(_apply_bold_reportlab(value), styles["Normal"]))
    doc.build(story)
    return buffer.getvalue()


def export_report(
    format: str,
    report_type: str,
    patient_name: str,
    report_date: date,
    content: str,
) -> tuple[bytes, str, str]:
    """Return (bytes, media_type, filename_suffix)."""
    if format in ("txt", "md"):
        exporter = export_txt if format == "txt" else export_md
        media = "text/plain; charset=utf-8" if format == "txt" else "text/markdown; charset=utf-8"
        return exporter(report_type, patient_name, report_date, content), media, format
    if format == "docx":
        return (
            export_docx(report_type, patient_name, report_date, content),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
    if format == "pdf":
        return (
            export_pdf(report_type, patient_name, report_date, content),
            "application/pdf",
            "pdf",
        )
    raise ValueError(f"Formato não suportado: {format}")
